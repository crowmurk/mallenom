from docx import Document
from docx.shared import Mm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl.utils.cell import get_column_letter

from django.http import HttpResponse
from django.utils.translation import gettext, gettext_lazy as _
from django.template.defaultfilters import date as date_filter

from .utils import DataBuilder


class ReportBuilder:
    empty_text = gettext("There are no records available")

    def __init__(self, start, end):
        """Создает отчеты в форматах docx, xlsx.

        Args:
            start - начальная дата отчета
            end - конечная дата отчетаа
        """

        self.start = start
        self.end = end
        self.data_builder = DataBuilder(self.start, self.end)

    def get_response(self, document, *, filename=None):
        """Возвращает объект http responce с отчетом.

        Args:
            document - документ с отчетом

        Kwargs:
            filename - имя файла отчета

        Returns: response
        """

        # Определяем дип документа
        try:
            document_type = document.__class__.__name__
        except AttributeError:
            document_type = None

        # Задаем тип контента и расширение файла
        if document_type == 'Document':
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = 'docx'
        elif document_type == 'Workbook':
            content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            file_extension = 'xlsx'
        else:
            raise ValueError("Invalid document type: {}".format(type(document)))

        # При необходимости, формируем имя файла
        if filename is None:
            filename = 'report_{}_{}'.format(self.start, self.end)
            filename = '.'.join([filename, file_extension])

        # Создаем ответ
        response = HttpResponse(content_type=content_type)
        response['Content-Disposition'] = 'attachment; filename={}'.format(filename)
        document.save(response)

        return response

    @staticmethod
    def _get_docx_document(*, orientation='portrait'):
        """Создает и настраивает пустой документ docx.

        Kwargs:
            orientation - ориентация документа

        Returns: docx документ
        """

        document = Document()
        section = document.sections[0]

        # Размер листа А4
        height = Mm(297)
        width = Mm(210)

        # Задаем ориентацию страницы
        if orientation == 'landscape':
            section.page_height = width
            section.page_width = height
        else:
            section.page_height = height
            section.page_width = width

        # Формат страницы
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        return document

    @staticmethod
    def _get_xlsx_document(*, orientation='portrait'):
        """Создает и настраивает пустой документ xlsx.

        Kwargs:
            orientation - ориентация документа

        Returns: xlsx документ
        """

        document = Workbook()
        sheet = document.active

        # Размер листа А4
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4

        # Ориентация страницы
        sheet.page_setup.orientation = sheet.ORIENTATION_PORTRAIT
        if orientation == 'landscape':
            sheet.page_setup.orientation = sheet.ORIENTATION_LANDSCAPE

        return document

    def _add_docx_title(self, document, title):
        """Добавляет заголовок в документы docx.

        Args:
            document - документ с отчетом
            title - название отчета

        Returns: None
        """

        heading = document.add_heading(title, level=1)
        heading.add_run().add_break()
        paragraph = document.add_paragraph(_('From '))
        paragraph.add_run(date_filter(self.start, "SHORT_DATE_FORMAT")).bold = True
        paragraph.add_run(_(' to '))
        paragraph.add_run(date_filter(self.end, "SHORT_DATE_FORMAT")).bold = True
        paragraph.add_run(':')
        paragraph.add_run().add_break()

    def _add_xlsx_title(self, sheet, title, width):
        """Добавляет заголовок в документы xlsx.

        Args:
            sheet - таблица с отчетом
            title - название отчета
            widw - число ячеек под заголовок

        Returns: None
        """

        sheet.insert_rows(1)
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=width)
        cell = sheet.cell(row=1, column=1)
        cell.value = title
        cell.font = Font(size=12, bold=True, italic=True)

        sheet.insert_rows(2)
        sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=width)

        sheet.insert_rows(3)
        sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=width)
        cell = sheet.cell(row=3, column=1)
        cell.value = gettext("From {start} to {end}:").format(
            start=date_filter(self.start, "SHORT_DATE_FORMAT"),
            end=date_filter(self.end, "SHORT_DATE_FORMAT"),
        )

        sheet.insert_rows(4)
        sheet.merge_cells(start_row=4, start_column=1, end_row=4, end_column=width)

    def assignment_report(self, **kwargs):
        """Формирует отчет.
        """

        # Получаем данные
        data = sorted(self.data_builder.assignment_report())

        # Создаем пустой документ
        document = self._get_docx_document(
            orientation=kwargs.get('orientation'),
        )
        self._add_docx_title(document, _("Employees' assignments"))

        if not data:
            # Возвращаем пустой отчет
            document.add_paragraph(self.empty_text)
            return document

        # Добавляем пустую таблицу
        table = document.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'

        # Заголовок таблицы
        header = (
            _('Employee'), _('Employee ID number'), _('Department'),
            _('Position'), _('Project'), _('Hours'),
        )

        # Добавляем заголовок в таблицу
        hdr_cells = table.rows[0].cells
        for col, value in enumerate(header):
            hdr_cells[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[col].paragraphs[0].add_run(value).bold = True

        # Добавляем данные в таблицу
        for row, row_value in enumerate(data):
            # Добавляем строку
            row_cells = table.add_row().cells
            # Проходим по ячейкам строки
            for col, cell_value in enumerate(row_value):
                # Объединяем ячейки Employee, Employee ID number
                if col in (0, 1) and str(cell_value) == table.cell(row, col).text:
                    # При совпадении с предыдущей строкой
                    table.cell(row, col).merge(row_cells[col])
                    continue
                # Объединяем ячейки Depatment, Position
                elif col in (2, 3) and table.cell(row, 1).text == str(row_value[1]):
                    # При совпадении с предыдущим ID Number
                    table.cell(row, col).merge(row_cells[col])
                    continue

                # Ячейки ID Number и Hours
                if col in (1, 5):
                    row_cells[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    row_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if col == 1:
                        # Добавляем значение ID Number в ячейку
                        row_cells[col].paragraphs[0].add_run(str(cell_value))
                    else:
                        # Значение Hours преобразуем в целое
                        # (при отсутствии дробной части)
                        if isinstance(cell_value, float) and cell_value.is_integer():
                            row_cells[col].paragraphs[0].add_run(
                                str(int(cell_value)),
                            )
                        else:
                            row_cells[col].paragraphs[0].add_run(
                                str(cell_value),
                            )
                    continue

                # Добавляем ячейки Employee, Department, Position, Project
                row_cells[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[col].paragraphs[0].add_run(cell_value)

        return document

    def assignment_matrix_report(self, **kwargs):
        """Формирует отчет.
        """

        absence_name = _('Absence hours')

        # Получаем данные
        data = self.data_builder.assignment_matrix_report(absence_name)

        # Создаем пустой документ
        document = self._get_docx_document(
            orientation=kwargs.get('orientation'),
        )
        self._add_docx_title(document, _("Employees' assignments matrix"))

        if not data:
            # Возвращаем пустой отчет
            document.add_paragraph(self.empty_text)
            return document

        # Список сотрудников с табельными номерами
        employes = tuple(sorted(set('{} [{}]'.format(
            *item[0:2]
        ) for item in data)))

        # Заголовок таблицы (проекты)
        header = sorted(list(set(
            item[2] for item in data if item[2] != absence_name
        )))
        header.insert(0, _('Employees'))
        header.append(absence_name)

        # Добавляем пустую таблицу
        table = document.add_table(rows=1, cols=len(header))
        table.style = 'Table Grid'

        # Добавляем заголовок в таблицу
        hdr_cells = table.rows[0].cells
        for index, value in enumerate(header):
            hdr_cell = hdr_cells[index]
            hdr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cell.paragraphs[0].add_run(value).bold = True

        # Добавляем вертикальный заголовок
        for value in employes:
            row_cells = table.add_row().cells
            row_cells[0].text = str(value)

        # Заполняем таблицу (матрицу)
        for employee, number, project, hours in data:
            # Определяем ячейку для записи
            row = employes.index('{} [{}]'.format(employee, number)) + 1
            col = header.index(project)

            # Добавляем данные в ячейку
            hours_cell = table.cell(row, col)
            hours_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hours_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Значение Hours преобразуем в целое
            # (при отсутствии дробной части)
            if isinstance(hours, float) and hours.is_integer():
                hours_cell.paragraphs[0].add_run(str(int(hours)))
            else:
                hours_cell.paragraphs[0].add_run(str(hours))

        return document

    def assignment_matrix_report_xlsx(self, **kwargs):
        """Формирует отчет.
        """

        title = gettext("Employees' assignments matrix")
        absence_name = gettext('Absence hours')

        # Получаем данные
        data = self.data_builder.assignment_matrix_report(absence_name)

        # Создаем пустой документ
        document = self._get_xlsx_document(
            orientation=kwargs.get('orientation'),
        )
        sheet = document.active
        sheet.title = title

        if not data:
            # Возвращаем пустой отчет
            self._add_xlsx_title(sheet, title, 9)
            sheet.merge_cells(start_row=5, start_column=1, end_row=5, end_column=9)
            cell = sheet.cell(row=5, column=1)
            cell.value = self.empty_text
            return document

        # Список сотрудников с табельными номерами
        employes = tuple(sorted(set('{} [{}]'.format(
            *item[0:2]
        ) for item in data)))

        # Заголовок таблицы (проекты)
        header = sorted(list(set(
            item[2] for item in data if item[2] != absence_name
        )))
        header.insert(0, gettext('Employees'))
        header.append(absence_name)

        # Добавяем заголовок в таблицу
        header_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')

        for col_num, column_title in enumerate(header, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = column_title
            cell.alignment = center_alignment
            cell.font = header_font

        # Добавляем вертикальный заголовок
        for row_num, value in enumerate(employes, 2):
            cell = sheet.cell(row=row_num, column=1)
            cell.value = value

        # Заполняем таблицу (матрицу)
        for employee, number, project, hours in data:
            row = employes.index('{} [{}]'.format(employee, number)) + 2
            col = header.index(project) + 1

            cell = sheet.cell(row=row, column=col)
            cell.value = hours
            cell.alignment = center_alignment

        # Устанавливаем ширину столбцов
        for column_cells in sheet.columns:
            length = max(len(str(cell.value)) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = length + 3

        # Границы таблицы
        border = Border(
            left=Side(border_style='thin', color='FF000000'),
            right=Side(border_style='thin', color='FF000000'),
            top=Side(border_style='thin', color='FF000000'),
            bottom=Side(border_style='thin', color='FF000000'),
        )
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = border

        # Добавляем заголовок отчета
        self._add_xlsx_title(sheet, title, len(header))

        return document

    def assignment_hours_check_xlsx(self, **kwargs):
        """Формирует отчет.
        """

        title = gettext("Employees' work hours check")

        # Получаем данные
        data = sorted(self.data_builder.assignment_hours_check())

        # Создаем пустой документ
        document = self._get_xlsx_document(
            orientation=kwargs.get('orientation'),
        )
        sheet = document.active
        sheet.title = title

        if not data:
            # Возвращаем пустой отчет
            self._add_xlsx_title(sheet, title, 10)
            sheet.merge_cells(start_row=5, start_column=1, end_row=5, end_column=10)
            cell = sheet.cell(row=5, column=1)
            cell.value = self.empty_text
            return document

        # Заголовок таблицы
        header = (
            gettext('Employee'), gettext('Employee ID number'),
            gettext('Department'), gettext('Position'), gettext('Staff units'),
            gettext('Hours assigned'), gettext('Absence hours'),
            gettext('Hours total'), gettext('Work hours'), gettext('Difference'),
        )
        header_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')

        # Добавляем заголовок в таблицу
        for col_num, column_title in enumerate(header, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = column_title
            cell.alignment = center_alignment
            cell.font = header_font

        # Добавляем строки в таблицу
        for row_num, item in enumerate(data, 2):
            for col_num, value in enumerate(item, 1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = value
                if col_num >= 5 or col_num == 2:
                    cell.alignment = center_alignment

        # Добавляем строку Total
        row_num = len(data) + 2
        cell = sheet.cell(row=row_num, column=1)
        cell.value = gettext('Total')
        cell.font = header_font
        for col_num, value in enumerate(header[4:], 5):
            cell = sheet.cell(row=row_num, column=col_num)
            cell.value = "=SUM({col}{start_row}:{col}{end_row})".format(
                col=get_column_letter(col_num),
                start_row=6,
                end_row=row_num + 3,
            )
            cell.alignment = center_alignment
            cell.font = header_font

        # Устанавливаем ширину столбцов
        for column_cells in sheet.columns:
            length = max(len(str(cell.value)) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = length + 3

        # Границы таблицы
        border = Border(
            left=Side(border_style='thin', color='FF000000'),
            right=Side(border_style='thin', color='FF000000'),
            top=Side(border_style='thin', color='FF000000'),
            bottom=Side(border_style='thin', color='FF000000'),
        )
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = border

        # Добавляем заголовок
        self._add_xlsx_title(sheet, title, len(header))

        # Объединяем ячейки заголовка строки Total
        sheet.merge_cells(start_row=row_num + 4, start_column=1, end_row=row_num + 4, end_column=4)

        return document

    def index_of_labor_distribution_xlsx(self, **kwargs):
        """Формирует отчет.
        """

        title = gettext("Employees' indexes of labor distribution")

        # Названия допролнительных полей
        absence_name = gettext('Absence hours')
        staff_units_name = gettext('Staff units')
        total_hours_name = gettext('Total')

        # Получаем данные
        data = self.data_builder.index_of_labor_distribution(absence_name)

        # Создаем пустой документ
        document = self._get_xlsx_document(
            orientation=kwargs.get('orientation'),
        )
        sheet = document.active
        sheet.title = title

        if not data:
            # Возвращаем пустой отчет
            self._add_xlsx_title(sheet, title, 9)
            sheet.merge_cells(start_row=5, start_column=1, end_row=5, end_column=9)
            cell = sheet.cell(row=5, column=1)
            cell.value = self.empty_text
            return document

        # Список сторудников с табельными номерами
        employes = tuple(sorted(set('{} [{}]'.format(
            *item[0:2]
        ) for item in data)))

        # Заголовок таблицы (проекты)
        header = sorted(list(set(
            item[3] for item in data if item[3] != absence_name
        )))
        header.insert(0, gettext('Employees'))
        header.extend((absence_name, total_hours_name, staff_units_name))

        header_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')

        # Добавляем заголовок в таблицу
        for col_num, column_title in enumerate(header, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = column_title
            cell.alignment = center_alignment
            cell.font = header_font

        # Добавляем сотрудников в таблицу
        for row_num, value in enumerate(employes, 2):
            cell = sheet.cell(row=row_num, column=1)
            cell.value = value

        # Добавляем данные в таблицу
        for employee, number, staff_units, project, hours, total_hours in data:
            row = employes.index('{} [{}]'.format(employee, number)) + 2
            col = header.index(project) + 1

            cell = sheet.cell(row=row, column=col)
            cell.value = hours
            cell.alignment = center_alignment

            col = header.index(staff_units_name) + 1
            cell = sheet.cell(row=row, column=col)
            cell.value = staff_units
            cell.alignment = center_alignment

            col = header.index(total_hours_name) + 1
            cell = sheet.cell(row=row, column=col)
            cell.value = total_hours
            cell.alignment = center_alignment

        # Добавляем строку Total
        row_num = len(employes) + 2
        cell = sheet.cell(row=row_num, column=1)
        cell.value = gettext('Total')
        cell.font = header_font
        for col_num, value in enumerate(header[1:], 2):
            cell = sheet.cell(row=row_num, column=col_num)
            cell.value = "=SUM({col}{start_row}:{col}{end_row})".format(
                col=get_column_letter(col_num),
                start_row=6,
                end_row=row_num + 3,
            )
            cell.alignment = center_alignment
            cell.font = header_font

        # Устанавливаем ширину столбцов
        for column_cells in sheet.columns:
            length = max(len(str(cell.value)) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = length + 3

        # Границы таблицы
        border = Border(
            left=Side(border_style='thin', color='FF000000'),
            right=Side(border_style='thin', color='FF000000'),
            top=Side(border_style='thin', color='FF000000'),
            bottom=Side(border_style='thin', color='FF000000'),
        )
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = border

        # Добавляем заголовок
        self._add_xlsx_title(sheet, title, len(header))

        return document

    def index_of_labor_distribution_per_project_xlsx(self, **kwargs):
        """Формирует отчет.
        """

        title = gettext("Employees' indexes of labor distribution per project")

        # Получаем данные
        data = self.data_builder.index_of_labor_distribution_per_project()

        # Создаем пустой документ
        document = self._get_xlsx_document(
            orientation=kwargs.get('orientation'),
        )
        sheet = document.active
        sheet.title = title

        if not data:
            # Возвращаем пустой отчет
            self._add_xlsx_title(sheet, title, 9)
            sheet.merge_cells(start_row=5, start_column=1, end_row=5, end_column=9)
            cell = sheet.cell(row=5, column=1)
            cell.value = self.empty_text
            return document

        # Список сторудников с табельными номерами
        employes = tuple(sorted(set('{} [{}]'.format(
            *item[0:2]
        ) for item in data)))

        # Создаем заголовок таблицы (проекты)
        header = sorted(list(set(item[2] for item in data)))
        header.insert(0, gettext('Employees'))

        header_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')

        # Добавляем заголовок в таблицу
        for col_num, column_title in enumerate(header, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = column_title
            cell.alignment = center_alignment
            cell.font = header_font

        # Добавляем сотрудников в таблицу
        for row_num, value in enumerate(employes, 2):
            cell = sheet.cell(row=row_num, column=1)
            cell.value = value

        # Добавляем данные в таблицу
        for employee, number, project, hours in data:
            row = employes.index('{} [{}]'.format(employee, number)) + 2
            col = header.index(project) + 1

            cell = sheet.cell(row=row, column=col)
            cell.value = hours
            cell.alignment = center_alignment

        # Добавляем строку Total
        row_num = len(employes) + 2
        cell = sheet.cell(row=row_num, column=1)
        cell.value = gettext('Total')
        cell.font = header_font
        for col_num, value in enumerate(header[1:], 2):
            cell = sheet.cell(row=row_num, column=col_num)
            cell.value = "=SUM({col}{start_row}:{col}{end_row})".format(
                col=get_column_letter(col_num),
                start_row=6,
                end_row=row_num + 3,
            )
            cell.alignment = center_alignment
            cell.font = header_font

        # Устанавливаем ширину столбцов
        for column_cells in sheet.columns:
            length = max(len(str(cell.value)) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = length + 3

        # Границы таблицы
        border = Border(
            left=Side(border_style='thin', color='FF000000'),
            right=Side(border_style='thin', color='FF000000'),
            top=Side(border_style='thin', color='FF000000'),
            bottom=Side(border_style='thin', color='FF000000'),
        )
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = border

        # Добавляем заголовок
        self._add_xlsx_title(sheet, title, len(header))

        return document
